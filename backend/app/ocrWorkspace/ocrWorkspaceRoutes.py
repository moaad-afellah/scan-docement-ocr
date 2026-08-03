import os
import io
import csv
import json
import uuid
import threading
from datetime import datetime
from pathlib import Path

from flask import Blueprint, request, jsonify, send_file, make_response, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename

from backend.extention import db
from backend.app.ocrWorkspace.ocrWorkspaceServices.ocrWorkspaceServicesProcess import (
    run_ocr,
    get_supported_engine_codes,
    is_supported_engine,
)
from backend.app.ocrWorkspace.ocrWorkspaceModels.ocrWorkspaceModelsUpload import (
    CANONICAL_DOCUMENT_TYPE_CODES,
    DocumentType,
    DocumentTypeField,
)
from backend.app.ocrWorkspace.ocrWorkspaceModels.ocrWorkspaceModelsProcess import Evaluation
from backend.app.ocrWorkspace.ocrWorkspaceModels.ocrWorkspaceModelsReview import EvaluationField
from backend.app.ocrWorkspace.ocrWorkspaceModels.ocrWorkspaceModelsBatch import BatchJob, BatchJobFile

# PDF export
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

# DOCX export
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

ocr_workspace_bp = Blueprint("ocr_workspace", __name__)

BACKEND_ROOT = Path(__file__).resolve().parents[2]
UPLOAD_FOLDER = BACKEND_ROOT / "uploads"

# Matches the "PNG, JPEG, WebP or PDF · up to 15 MB per file" copy in the
# upload dropzone -- this used to be UI-only; now it's actually enforced.
ALLOWED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".pdf"}
MAX_FILE_SIZE_BYTES = 15 * 1024 * 1024  # 15 MB


# ---------------------------------------------------------------------------
# Shared upload helper
# ---------------------------------------------------------------------------

def save_upload_file(file_storage):
    """
    Validates and safely persists an uploaded file.
    Returns (file_path, safe_display_name).
    Raises ValueError with a user-facing message on any validation failure.
    """
    if not file_storage or file_storage.filename == "":
        raise ValueError("No file selected")

    ext = os.path.splitext(file_storage.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Allowed types: "
            f"{', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    # Check size without loading the whole file into memory
    file_storage.stream.seek(0, os.SEEK_END)
    size = file_storage.stream.tell()
    file_storage.stream.seek(0)
    if size > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File '{file_storage.filename}' exceeds the 15 MB limit"
        )

    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

    safe_name = secure_filename(file_storage.filename)
    if not safe_name:
        safe_name = f"upload{ext}"

    # Prefix with a random token so two users (or the same user twice)
    # uploading "invoice.pdf" can never collide or overwrite each other.
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"
    file_path = str(UPLOAD_FOLDER / unique_name)
    file_storage.save(file_path)

    return file_path, safe_name


# ---------------------------------------------------------------------------
# Test OCR (dev/debug endpoint) -- now requires auth like everything else
# ---------------------------------------------------------------------------

@ocr_workspace_bp.route("/test-ocr", methods=["POST"])
@jwt_required()
def test_ocr():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    engine = request.form.get("engine", "pytesseract")

    try:
        file_path, _safe_name = save_upload_file(file)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    try:
        result = run_ocr(file_path, engine=engine)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    return jsonify(result), 200


# ---------------------------------------------------------------------------
# Document types
# ---------------------------------------------------------------------------

@ocr_workspace_bp.route("/document-types", methods=["GET"])
def list_document_types():
    allowed_codes = set(CANONICAL_DOCUMENT_TYPE_CODES)
    doc_types = DocumentType.query.filter(DocumentType.code.in_(allowed_codes)).all()
    ordered_doc_types = sorted(
        doc_types,
        key=lambda dt: CANONICAL_DOCUMENT_TYPE_CODES.index(dt.code) if dt.code in CANONICAL_DOCUMENT_TYPE_CODES else len(CANONICAL_DOCUMENT_TYPE_CODES),
    )

    return jsonify([
        {
            "id": dt.id,
            "code": dt.code,
            "name": dt.name,
            "fields": [
                {"id": f.id, "label": f.label, "position": f.position}
                for f in sorted(dt.fields, key=lambda f: f.position)
            ]
        }
        for dt in ordered_doc_types
    ]), 200


@ocr_workspace_bp.route("/document-types", methods=["POST"])
@jwt_required()
def create_document_type():
    data = request.get_json()
    if not data:
        return jsonify({"error": "No JSON data provided"}), 400

    code = data.get("code")
    name = data.get("name")
    if not code or not name:
        return jsonify({"error": "code and name are required"}), 400

    existing = DocumentType.query.filter_by(code=code).first()
    if existing:
        return jsonify({"error": "A document type with this code already exists"}), 409

    new_doc_type = DocumentType(code=code, name=name)
    db.session.add(new_doc_type)
    db.session.commit()

    return jsonify({"id": new_doc_type.id, "code": new_doc_type.code, "name": new_doc_type.name, "fields": []}), 201


@ocr_workspace_bp.route("/document-types/<int:doc_type_id>/fields", methods=["POST"])
@jwt_required()
def add_document_type_field(doc_type_id):
    doc_type = DocumentType.query.get(doc_type_id)
    if not doc_type:
        return jsonify({"error": "Document type not found"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"error": "No JSON data provided"}), 400

    fields_to_create = data if isinstance(data, list) else [data]

    created_fields = []
    for field_data in fields_to_create:
        label = field_data.get("label")
        position = field_data.get("position")
        if not label or position is None:
            return jsonify({"error": "label and position are required for every field"}), 400

        new_field = DocumentTypeField(document_type_id=doc_type_id, label=label, position=position)
        db.session.add(new_field)
        created_fields.append(new_field)

    db.session.commit()

    return jsonify([
        {
            "id": f.id,
            "document_type_id": f.document_type_id,
            "label": f.label,
            "position": f.position
        }
        for f in created_fields
    ]), 201


# ---------------------------------------------------------------------------
# Single evaluation (synchronous -- fine for one file)
# ---------------------------------------------------------------------------

@ocr_workspace_bp.route("/evaluations", methods=["POST"])
@jwt_required()
def create_evaluation():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    engine_id = request.form.get("engine_id")
    document_type_id = request.form.get("document_type_id")
    user_id = int(get_jwt_identity())

    if not engine_id or not document_type_id:
        return jsonify({"error": "engine_id and document_type_id are required"}), 400

    try:
        engine_id = int(engine_id)
        document_type_id = int(document_type_id)
    except ValueError:
        return jsonify({"error": "engine_id and document_type_id must be integers"}), 400

    doc_type = DocumentType.query.get(document_type_id)
    if not doc_type:
        return jsonify({"error": "Document type not found"}), 404

    from backend.app.settings.settingsModels import OcrEngine
    engine = OcrEngine.query.get(engine_id)
    if not engine:
        return jsonify({"error": "OCR engine not found"}), 404

    if not is_supported_engine(engine.code):
        supported = ", ".join(get_supported_engine_codes())
        return jsonify({"error": f"Unsupported OCR engine '{engine.code}'. Supported engines: {supported}"}), 400

    try:
        file_path, safe_name = save_upload_file(file)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    file_size = os.path.getsize(file_path)

    new_evaluation = Evaluation(
        user_id=user_id,
        document_type_id=document_type_id,
        engine_id=engine_id,
        file_name=safe_name,
        file_size=file_size,
        file_path=file_path,
        status='pending'
    )
    db.session.add(new_evaluation)
    db.session.flush()

    try:
        sorted_fields = sorted(doc_type.fields, key=lambda f: f.position)
        field_labels = [f.label for f in sorted_fields]

        ocr_result = run_ocr(file_path, engine=engine.code, field_labels=field_labels)
        extracted_fields = ocr_result.get("fields", {})
        new_evaluation.raw_text = ocr_result.get("raw_text")

        for field_def in sorted_fields:
            key = field_def.label.lower().replace(" ", "_")
            ocr_value = extracted_fields.get(key)

            eval_field = EvaluationField(
                evaluation_id=new_evaluation.id,
                label=field_def.label,
                ocr_value=ocr_value,
                reference_value=None,
                status='pending',
                position=field_def.position
            )
            db.session.add(eval_field)

        new_evaluation.status = 'done'

    except Exception as e:
        current_app.logger.exception("OCR processing failed for evaluation %s", new_evaluation.id)
        new_evaluation.status = 'error'
        new_evaluation.error_message = str(e) if str(e) else "OCR processing failed. Please try again or contact support."

    db.session.commit()

    return jsonify({
        "id": new_evaluation.id,
        "status": new_evaluation.status,
        "raw_text": new_evaluation.raw_text,
        "error_message": new_evaluation.error_message,
        "fields": [
            {
                "id": f.id,
                "label": f.label,
                "ocr_value": f.ocr_value,
                "position": f.position
            }
            for f in sorted(new_evaluation.fields, key=lambda f: f.position)
        ]
    }), 201


@ocr_workspace_bp.route("/evaluations/<int:evaluation_id>/fields", methods=["PUT"])
@jwt_required()
def submit_reference_values(evaluation_id):
    evaluation = Evaluation.query.get(evaluation_id)
    if not evaluation:
        return jsonify({"error": "Evaluation not found"}), 404

    current_user_id = int(get_jwt_identity())
    if evaluation.user_id != current_user_id:
        return jsonify({"error": "Unauthorized to modify this evaluation"}), 403

    data = request.get_json()
    if not data or "fields" not in data:
        return jsonify({"error": "fields data is required"}), 400

    existing_field_ids = {str(f.id): f for f in evaluation.fields}

    for key, field_data in data["fields"].items():
        if key in existing_field_ids:
            eval_field = existing_field_ids[key]
            eval_field.reference_value = field_data.get("reference_value")
            eval_field.status = field_data.get("status")
            if "label" in field_data:
                eval_field.label = field_data["label"]
        else:
            label = field_data.get("label", key)
            max_pos = max([f.position for f in evaluation.fields], default=0)
            new_field = EvaluationField(
                evaluation_id=evaluation.id,
                label=label,
                ocr_value=field_data.get("ocr_value"),
                reference_value=field_data.get("reference_value"),
                status=field_data.get("status", "pending"),
                position=max_pos + 1
            )
            db.session.add(new_field)

    if "new_fields" in data and isinstance(data["new_fields"], list):
        max_pos = max([f.position for f in evaluation.fields], default=0)
        for idx, new_f in enumerate(data["new_fields"]):
            new_field = EvaluationField(
                evaluation_id=evaluation.id,
                label=new_f.get("label", "Custom Field"),
                ocr_value=new_f.get("ocr_value"),
                reference_value=new_f.get("reference_value"),
                status=new_f.get("status", "pending"),
                position=max_pos + idx + 1
            )
            db.session.add(new_field)

    db.session.flush()

    # --- Server-side accuracy computation -------------------------------
    # Previously the client computed correct/incorrect/missing/additional
    # counts and accuracy itself, and the server just stored whatever it
    # was told -- trivially gameable. Now we recompute from the actual
    # ocr_value / reference_value pairs so the dashboard leaderboard means
    # something. The client can still send counts for a "preview" but they
    # are ignored in favor of this calculation.
    correct = incorrect = missing = additional = 0
    for f in evaluation.fields:
        ocr_val = (f.ocr_value or "").strip()
        ref_val = (f.reference_value or "").strip()

        if not ref_val:
            continue  # reviewer hasn't entered a reference value yet -- skip

        if not ocr_val and ref_val:
            missing += 1
            f.status = "missing"
        elif ocr_val and not ref_val:
            additional += 1
            f.status = "additional"
        elif ocr_val.lower() == ref_val.lower():
            correct += 1
            f.status = "match"
        else:
            incorrect += 1
            f.status = "mismatch"

    total = correct + incorrect + missing + additional
    accuracy = round((correct / total) * 100) if total else 0

    evaluation.correct_count = correct
    evaluation.incorrect_count = incorrect
    evaluation.missing_count = missing
    evaluation.additional_count = additional
    evaluation.total_count = total
    evaluation.accuracy = accuracy

    db.session.commit()

    return jsonify({
        "id": evaluation.id,
        "accuracy": evaluation.accuracy,
        "correct_count": evaluation.correct_count,
        "incorrect_count": evaluation.incorrect_count,
        "missing_count": evaluation.missing_count,
        "additional_count": evaluation.additional_count,
        "total_count": evaluation.total_count,
        "fields": [
            {
                "id": f.id,
                "label": f.label,
                "ocr_value": f.ocr_value,
                "reference_value": f.reference_value,
                "status": f.status,
                "position": f.position
            }
            for f in sorted(evaluation.fields, key=lambda f: f.position)
        ]
    }), 200


@ocr_workspace_bp.route("/evaluations/<int:evaluation_id>", methods=["GET"])
@jwt_required()
def get_evaluation(evaluation_id):
    evaluation = Evaluation.query.get(evaluation_id)
    if not evaluation:
        return jsonify({"error": "Evaluation not found"}), 404

    current_user_id = int(get_jwt_identity())
    if evaluation.user_id != current_user_id:
        return jsonify({"error": "Unauthorized to view this evaluation"}), 403

    return jsonify({
        "id": evaluation.id,
        "user_id": evaluation.user_id,
        "document_type_id": evaluation.document_type_id,
        "engine_id": evaluation.engine_id,
        "file_name": evaluation.file_name,
        "status": evaluation.status,
        "raw_text": evaluation.raw_text,
        "accuracy": evaluation.accuracy,
        "correct_count": evaluation.correct_count,
        "incorrect_count": evaluation.incorrect_count,
        "missing_count": evaluation.missing_count,
        "additional_count": evaluation.additional_count,
        "total_count": evaluation.total_count,
        "created_at": evaluation.created_at.isoformat(),
        "fields": [
            {
                "id": f.id,
                "label": f.label,
                "ocr_value": f.ocr_value,
                "reference_value": f.reference_value,
                "status": f.status,
                "position": f.position
            }
            for f in sorted(evaluation.fields, key=lambda f: f.position)
        ]
    }), 200


@ocr_workspace_bp.route("/evaluations", methods=["GET"])
@jwt_required()
def list_evaluations():
    current_user_id = int(get_jwt_identity())
    evaluations = Evaluation.query.filter_by(user_id=current_user_id).order_by(Evaluation.created_at.desc()).all()

    return jsonify([
        {
            "id": e.id,
            "file_name": e.file_name,
            "status": e.status,
            "accuracy": e.accuracy,
            "created_at": e.created_at.isoformat()
        }
        for e in evaluations
    ]), 200


@ocr_workspace_bp.route("/evaluations/<int:evaluation_id>/file", methods=["GET"])
@jwt_required()
def get_evaluation_file(evaluation_id):
    evaluation = Evaluation.query.get(evaluation_id)
    if not evaluation:
        return jsonify({"error": "Evaluation not found"}), 404

    current_user_id = int(get_jwt_identity())
    if evaluation.user_id != current_user_id:
        return jsonify({"error": "Unauthorized to access this file"}), 403

    if not evaluation.file_path or not os.path.exists(evaluation.file_path):
        return jsonify({"error": "File not found on server"}), 404

    return send_file(evaluation.file_path)


# ---------------------------------------------------------------------------
# Batch evaluations -- now genuinely async with a pollable job status
# ---------------------------------------------------------------------------

def _process_batch_job(app, job_id):
    """
    Runs in a background thread. Processes each file in the batch job
    sequentially, committing after every file so GET /evaluations/batch/<id>
    reflects live progress (matches the "1/1 completed" checklist UI).
    """
    with app.app_context():
        job = BatchJob.query.get(job_id)
        if not job:
            return

        job.status = 'processing'
        db.session.commit()

        doc_type = DocumentType.query.get(job.document_type_id)
        from backend.app.settings.settingsModels import OcrEngine
        engine = OcrEngine.query.get(job.engine_id)

        sorted_fields = sorted(doc_type.fields, key=lambda f: f.position) if doc_type else []
        field_labels = [f.label for f in sorted_fields]

        job_files = BatchJobFile.query.filter_by(batch_job_id=job.id).order_by(BatchJobFile.position).all()

        for jf in job_files:
            jf.status = 'processing'
            db.session.commit()

            try:
                file_size = os.path.getsize(jf.file_path)

                new_evaluation = Evaluation(
                    user_id=job.user_id,
                    document_type_id=job.document_type_id,
                    engine_id=job.engine_id,
                    file_name=jf.file_name,
                    file_size=file_size,
                    file_path=jf.file_path,
                    status='pending'
                )
                db.session.add(new_evaluation)
                db.session.flush()

                ocr_result = run_ocr(jf.file_path, engine=engine.code, field_labels=field_labels)
                extracted_fields = ocr_result.get("fields", {})
                new_evaluation.raw_text = ocr_result.get("raw_text")

                for field_def in sorted_fields:
                    key = field_def.label.lower().replace(" ", "_")
                    ocr_value = extracted_fields.get(key)
                    db.session.add(EvaluationField(
                        evaluation_id=new_evaluation.id,
                        label=field_def.label,
                        ocr_value=ocr_value,
                        reference_value=None,
                        status='pending',
                        position=field_def.position
                    ))

                new_evaluation.status = 'done'
                jf.evaluation_id = new_evaluation.id
                jf.status = 'done'

            except Exception as e:
                app.logger.exception("Batch OCR failed for file %s in job %s", jf.file_name, job.id)
                jf.status = 'error'
                jf.error_message = str(e) if str(e) else "OCR processing failed for this file."

            job.completed_files += 1
            db.session.commit()

        job.status = 'completed'
        job.completed_at = datetime.utcnow()
        db.session.commit()


@ocr_workspace_bp.route("/evaluations/batch", methods=["POST"])
@jwt_required()
def create_evaluations_batch():
    files = request.files.getlist("files")
    if not files or all(f.filename == "" for f in files):
        return jsonify({"error": "No files provided"}), 400

    engine_id = request.form.get("engine_id")
    document_type_id = request.form.get("document_type_id")
    user_id = int(get_jwt_identity())

    if not engine_id or not document_type_id:
        return jsonify({"error": "engine_id and document_type_id are required"}), 400

    try:
        engine_id = int(engine_id)
        document_type_id = int(document_type_id)
    except ValueError:
        return jsonify({"error": "engine_id and document_type_id must be integers"}), 400

    doc_type = DocumentType.query.get(document_type_id)
    if not doc_type:
        return jsonify({"error": "Document type not found"}), 404

    from backend.app.settings.settingsModels import OcrEngine
    engine = OcrEngine.query.get(engine_id)
    if not engine:
        return jsonify({"error": "OCR engine not found"}), 404

    if not is_supported_engine(engine.code):
        supported = ", ".join(get_supported_engine_codes())
        return jsonify({"error": f"Unsupported OCR engine '{engine.code}'. Supported engines: {supported}"}), 400

    job = BatchJob(
        user_id=user_id,
        document_type_id=document_type_id,
        engine_id=engine_id,
        status='pending'
    )
    db.session.add(job)
    db.session.flush()

    saved_count = 0
    for idx, file in enumerate(files):
        if file.filename == "":
            continue
        try:
            file_path, safe_name = save_upload_file(file)
        except ValueError as e:
            db.session.rollback()
            return jsonify({"error": str(e)}), 400

        db.session.add(BatchJobFile(
            batch_job_id=job.id,
            file_name=safe_name,
            file_path=file_path,
            status='pending',
            position=idx
        ))
        saved_count += 1

    if saved_count == 0:
        db.session.rollback()
        return jsonify({"error": "No valid files selected"}), 400

    job.total_files = saved_count
    db.session.commit()

    # Kick off background processing and return immediately so the frontend
    # can start polling GET /evaluations/batch/<job_id> for progress.
    app_obj = current_app._get_current_object()
    thread = threading.Thread(target=_process_batch_job, args=(app_obj, job.id), daemon=True)
    thread.start()

    return jsonify({
        "job_id": job.id,
        "status": job.status,
        "total_files": job.total_files,
        "completed_files": job.completed_files
    }), 202


@ocr_workspace_bp.route("/evaluations/batch/<job_id>", methods=["GET"])
@jwt_required()
def get_batch_job_status(job_id):
    job = BatchJob.query.get(job_id)
    if not job:
        return jsonify({"error": "Batch job not found"}), 404

    current_user_id = int(get_jwt_identity())
    if job.user_id != current_user_id:
        return jsonify({"error": "Unauthorized to view this batch job"}), 403

    files = BatchJobFile.query.filter_by(batch_job_id=job.id).order_by(BatchJobFile.position).all()

    return jsonify({
        "job_id": job.id,
        "status": job.status,
        "total_files": job.total_files,
        "completed_files": job.completed_files,
        "created_at": job.created_at.isoformat() if job.created_at else None,
        "completed_at": job.completed_at.isoformat() if job.completed_at else None,
        "files": [
            {
                "id": jf.id,
                "file_name": jf.file_name,
                "status": jf.status,
                "evaluation_id": jf.evaluation_id,
                "error_message": jf.error_message
            }
            for jf in files
        ]
    }), 200


# ---------------------------------------------------------------------------
# Export -- now actually produces csv / txt / pdf / docx / json
# ---------------------------------------------------------------------------

def _build_export_pdf(evaluation, sorted_fields):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=18, spaceAfter=4)
    meta_style = ParagraphStyle('MetaStyle', parent=styles['Normal'], fontSize=10, textColor=colors.grey)

    elements = []
    elements.append(Paragraph("VeraScan Evaluation Export", title_style))
    elements.append(Paragraph(f"File: {evaluation.file_name}", meta_style))
    elements.append(Paragraph(f"Engine: {evaluation.engine.name if evaluation.engine else 'N/A'}", meta_style))
    elements.append(Paragraph(
        f"Document Type: {evaluation.document_type.name if evaluation.document_type else 'N/A'}", meta_style
    ))
    elements.append(Paragraph(f"Accuracy: {evaluation.accuracy or 0}%", meta_style))
    elements.append(Paragraph(
        f"Status Counts: {evaluation.correct_count or 0} Correct, "
        f"{evaluation.incorrect_count or 0} Incorrect, "
        f"{evaluation.missing_count or 0} Missing, "
        f"{evaluation.additional_count or 0} Additional",
        meta_style
    ))
    if evaluation.created_at:
        elements.append(Paragraph(f"Date: {evaluation.created_at.isoformat()}", meta_style))
    elements.append(Spacer(1, 0.25 * inch))

    table_data = [["Field", "OCR Value", "Reference Value", "Status"]]
    for f in sorted_fields:
        table_data.append([f.label, f.ocr_value or "-", f.reference_value or "-", (f.status or "pending").title()])

    table = Table(table_data, colWidths=[1.6 * inch, 1.8 * inch, 1.8 * inch, 1.2 * inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#f3f4f6")]),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(table)

    doc.build(elements)
    buffer.seek(0)
    return buffer


def _build_export_docx(evaluation, sorted_fields):
    document = DocxDocument()

    heading = document.add_heading("VeraScan Evaluation Export", level=1)

    meta_lines = [
        f"File: {evaluation.file_name}",
        f"Engine: {evaluation.engine.name if evaluation.engine else 'N/A'}",
        f"Document Type: {evaluation.document_type.name if evaluation.document_type else 'N/A'}",
        f"Accuracy: {evaluation.accuracy or 0}%",
        f"Status Counts: {evaluation.correct_count or 0} Correct, "
        f"{evaluation.incorrect_count or 0} Incorrect, "
        f"{evaluation.missing_count or 0} Missing, "
        f"{evaluation.additional_count or 0} Additional",
    ]
    if evaluation.created_at:
        meta_lines.append(f"Date: {evaluation.created_at.isoformat()}")

    for line in meta_lines:
        p = document.add_paragraph(line)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    document.add_paragraph("")

    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(["Field", "OCR Value", "Reference Value", "Status"]):
        hdr_cells[i].text = header

    for f in sorted_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = f.label
        row_cells[1].text = f.ocr_value or "-"
        row_cells[2].text = f.reference_value or "-"
        row_cells[3].text = (f.status or "pending").title()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@ocr_workspace_bp.route("/evaluations/<int:evaluation_id>/export", methods=["GET"])
@jwt_required()
def export_evaluation(evaluation_id):
    evaluation = Evaluation.query.get(evaluation_id)
    if not evaluation:
        return jsonify({"error": "Evaluation not found"}), 404

    current_user_id = int(get_jwt_identity())
    if evaluation.user_id != current_user_id:
        return jsonify({"error": "Unauthorized to export this evaluation"}), 403

    export_format = request.args.get("format", "json").lower()
    file_basename = os.path.splitext(evaluation.file_name)[0]
    sorted_fields = sorted(evaluation.fields, key=lambda f: f.position)

    if export_format == "csv":
        si = io.StringIO()
        writer = csv.writer(si)
        writer.writerow(["Field Label", "OCR Value", "Reference Value", "Status"])
        for f in sorted_fields:
            writer.writerow([f.label, f.ocr_value or "", f.reference_value or "", f.status or ""])

        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename={file_basename}_evaluation.csv"
        output.headers["Content-type"] = "text/csv"
        return output

    elif export_format in ["txt", "text"]:
        lines = [
            "VeraScan Evaluation Export",
            "==========================",
            f"File Name: {evaluation.file_name}",
            f"Engine: {evaluation.engine.name if evaluation.engine else 'N/A'}",
            f"Document Type: {evaluation.document_type.name if evaluation.document_type else 'N/A'}",
            f"Accuracy: {evaluation.accuracy or 0}%",
            f"Status Counts: {evaluation.correct_count or 0} Correct, "
            f"{evaluation.incorrect_count or 0} Incorrect, "
            f"{evaluation.missing_count or 0} Missing, "
            f"{evaluation.additional_count or 0} Additional",
            f"Date: {evaluation.created_at.isoformat() if evaluation.created_at else ''}",
            "",
            "FIELD COMPARISON:",
            "-----------------"
        ]
        for f in sorted_fields:
            lines.append(f"Label: {f.label}")
            lines.append(f"  OCR Output : {f.ocr_value or '-'}")
            lines.append(f"  Reference  : {f.reference_value or '-'}")
            lines.append(f"  Status     : {f.status or 'Pending'}\n")

        output = make_response("\n".join(lines))
        output.headers["Content-Disposition"] = f"attachment; filename={file_basename}_evaluation.txt"
        output.headers["Content-type"] = "text/plain"
        return output

    elif export_format == "pdf":
        buffer = _build_export_pdf(evaluation, sorted_fields)
        return send_file(
            buffer,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{file_basename}_evaluation.pdf"
        )

    elif export_format == "docx":
        buffer = _build_export_docx(evaluation, sorted_fields)
        return send_file(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{file_basename}_evaluation.docx"
        )

    else:  # json / default
        export_data = {
            "evaluation_id": evaluation.id,
            "file_name": evaluation.file_name,
            "document_type": evaluation.document_type.name if evaluation.document_type else None,
            "engine": evaluation.engine.name if evaluation.engine else None,
            "accuracy": evaluation.accuracy,
            "summary": {
                "correct_count": evaluation.correct_count,
                "incorrect_count": evaluation.incorrect_count,
                "missing_count": evaluation.missing_count,
                "additional_count": evaluation.additional_count,
                "total_count": evaluation.total_count
            },
            "created_at": evaluation.created_at.isoformat() if evaluation.created_at else None,
            "fields": [
                {
                    "id": f.id,
                    "label": f.label,
                    "ocr_value": f.ocr_value,
                    "reference_value": f.reference_value,
                    "status": f.status,
                    "position": f.position
                }
                for f in sorted_fields
            ]
        }
        output = make_response(json.dumps(export_data, indent=2))
        output.headers["Content-Disposition"] = f"attachment; filename={file_basename}_evaluation.json"
        output.headers["Content-type"] = "application/json"
        return output